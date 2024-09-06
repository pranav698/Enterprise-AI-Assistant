import streamlit as st
import hashlib
import sqlite3
import smtplib
import random
import re
import pandas as pd
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
import os
import tempfile
from ragpart import generate_response_from_chunks, get_relevant_chunks, create_index, extract_text_from_pdf, clean_text, store_chunks_in_pinecone, combined_chunking
from translate import translate, generate_audio

# Initialize session state
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False

if 'queries_and_responses' not in st.session_state:
    st.session_state.queries_and_responses = []

# Connect to SQLite Database
conn = sqlite3.connect('users.db')
c = conn.cursor()

# Create users table
c.execute('''
CREATE TABLE IF NOT EXISTS users (
    email TEXT PRIMARY KEY,
    password TEXT
)
''')
conn.commit()

# Function to hash passwords
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Function to check if the user exists
def check_user(email):
    c.execute('SELECT * FROM users WHERE email = ?', (email,))
    return c.fetchone()

# Function to add a new user
def add_user(email, password):
    hashed_password = hash_password(password)
    c.execute('INSERT INTO users (email, password) VALUES (?, ?)', 
              (email, hashed_password))
    conn.commit()

# Function to generate a random 6-digit OTP
def generate_otp():
    return str(random.randint(100000, 999999))

def process_local_pdfs(data):
    combined_chunks = []
    
    # Check if data is a DataFrame
    if isinstance(data, pd.DataFrame):
        data = data.to_dict()
        data = data['text']

    # If data is a list of uploaded files
    for pdf_file in data:
        if isinstance(data, dict) and isinstance(data[pdf_file], str):
            text = data[pdf_file]  
        else:
            text = extract_text_from_pdf(pdf_file)
        
        cleaned_text = clean_text(text)
        chunks = combined_chunking(cleaned_text)
        combined_chunks.extend(chunks)
    
    return combined_chunks


# Function to send OTP via email using Outlook SMTP server
def send_otp_via_email(email, otp):
    sender_email = 'sihbrocode@outlook.com'  # Replace with your email
    sender_password = 'Rohith@123'  # Replace with your email password

    subject = "Your OTP for Login"
    body = f"Your OTP is: {otp}"

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    try:
        # Connect to the Outlook SMTP server
        with smtplib.SMTP('smtp.office365.com', 587) as server:
            server.starttls()  # Upgrade to a secure connection
            server.login(sender_email, sender_password)  # Login to the email server
            server.send_message(msg)  # Send the email
        st.success("OTP sent successfully to your email.")
    except Exception as e:
        st.error(f"Failed to send OTP: {e}")

# Function to send the Word document via email
def send_word_document_via_email(email, queries_and_responses):
    sender_email = 'sihbrocode@outlook.com'
    sender_password = 'Rohith@123'

    subject = "Your Queries and Responses"
    body = "Attached is the Word document containing your queries and responses."

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = email
    msg['Subject'] = subject

    msg.attach(MIMEText(body, 'plain'))

    # Create Word document
    doc = Document()
    for entry in queries_and_responses:
        doc.add_heading('Query:', level=2)
        doc.add_paragraph(entry['query'])
        doc.add_heading('Response:', level=2)
        doc.add_paragraph(entry['response'])
        doc.add_paragraph('---')

    # Save document to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        doc.save(temp_file.name)
        temp_file.seek(0)
        
        # Attach the Word document
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(temp_file.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename=queries_and_responses.docx')
        msg.attach(part)

    try:
        with smtplib.SMTP('smtp.office365.com', 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
        st.success("Document sent successfully to your email.")
    except Exception as e:
        st.error(f"Failed to send the document: {e}")
    finally:
        os.unlink(temp_file.name)

# Function to update the Word document with queries and responses
def update_word_document(document_path, queries_and_responses):
    doc = Document()

    for entry in queries_and_responses:
        doc.add_heading('Query:', level=2)
        doc.add_paragraph(entry['query'])
        doc.add_heading('Response:', level=2)
        doc.add_paragraph(entry['response'])
        doc.add_paragraph('---')

    doc.save(document_path)

def handle_query_response(query, lang):
    relevant_chunks = get_relevant_chunks(query, st.session_state.index)
    response = generate_response_from_chunks(relevant_chunks, query)

    # Store query and response in session state
    if 'queries_and_responses' not in st.session_state:
        st.session_state.queries_and_responses = []
    st.session_state.queries_and_responses.append({
        "query": query,
        "response": response
    })

    # Display the response
    if lang != "English":
        translated_response = translate(response, lang)
        st.write(translated_response)
        audio_io = generate_audio(translated_response, lang)
    else:
        st.write(response)
        audio_io = generate_audio(response, lang)

    st.audio(audio_io, format='audio/mp3')


def load_bad_words(file_path='bad_words.txt'):
    with open(file_path, 'r') as file:
        return [word.strip().lower() for word in file.readlines()]
    

def contains_bad_words(text, bad_words):
    return any(word.lower() in text.lower() for word in bad_words)

BAD_WORDS = load_bad_words()

def main_app():
    st.sidebar.image("logo.jpg")
    st.title("Enterprise AI Assistant")
    st.sidebar.title("Enterprise AI Assistant")

    lang = st.sidebar.radio("Choose", ["English", "French", "Spanish"])

    # Handle Local PDF Processing
    data = st.sidebar.file_uploader("Upload a PDF", type="pdf", accept_multiple_files=True)
    if data and 'papers_downloaded' not in st.session_state:
        with st.spinner("Processing PDFs..."):
            combined_chunks = process_local_pdfs(data)
            st.session_state.index = create_index()
            if st.session_state.index:
                store_chunks_in_pinecone(combined_chunks, st.session_state.index)
                st.session_state.papers_downloaded = True
                st.success("PDF processed and indexed successfully!")
            else:
                st.error("Failed to create Pinecone index.")

    # Query handling
    if 'index' in st.session_state:
        query = st.text_input("Enter your question:")
        if query:
            # Check for bad words
            if contains_bad_words(query, BAD_WORDS):
                st.warning("Your message contains inappropriate language. Please rephrase your question.")
            else:
                if st.button("Ask", key=f"ask_button_{hash(query)}"):
                    with st.spinner("Searching for answers..."):
                        handle_query_response(query, lang)

        # Add button to send queries and responses as a Word document
        if st.button("Send Queries and Responses as Word Document", key=f"send_document_button_{hash(str(st.session_state.get('queries_and_responses', '')))}"):
            if 'queries_and_responses' in st.session_state and st.session_state.queries_and_responses:
                send_word_document_via_email(st.session_state.email, st.session_state.queries_and_responses)
            else:
                st.warning("No queries and responses to send.")

        if st.button("End conversation", key="end_conversation_button"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.experimental_rerun()

def validate_password(password):
    rules = {
        "min_length": len(password) >= 8,
        "uppercase": bool(re.search(r"[A-Z]", password)),
        "lowercase": bool(re.search(r"[a-z]", password)),
        "digit": bool(re.search(r"[0-9]", password)),
        "special_char": bool(re.search(r"[@#$%^&+=]", password))
    }
    return rules

def login_page():
    st.title('Two-Factor Authentication App')

    # Option to login or register
    option = st.selectbox('Choose an option:', ['Login', 'Register'])

    if option == 'Register':
        st.subheader('Register New User')
        email = st.text_input('Email')
        password = st.text_input('Password', type='password')

        st.write("**Password Strength Rules:**")
        password_rules = validate_password(password)
        st.write(f"- **At least 8 characters long**: {'✅' if password_rules['min_length'] else '❌'}")
        st.write(f"- **At least one uppercase letter**: {'✅' if password_rules['uppercase'] else '❌'}")
        st.write(f"- **At least one lowercase letter**: {'✅' if password_rules['lowercase'] else '❌'}")
        st.write(f"- **At least one digit**: {'✅' if password_rules['digit'] else '❌'}")
        st.write(f"- **At least one special character**: {'✅' if password_rules['special_char'] else '❌'}")

        if st.button('Register'):
            if check_user(email):
                st.error('User already exists. Please login.')
            else:
                if all(password_rules.values()):
                    add_user(email, password)
                    st.success('Registration successful! Please log in.')
                else:
                    st.error('Password does not meet the strength requirements.')

    elif option == 'Login':
        st.subheader('User Login')
        email = st.text_input('Email')
        password = st.text_input('Password', type='password')

        if st.button('Login'):
            user = check_user(email)
            if user and user[1] == hash_password(password):
                # Generate OTP
                otp = generate_otp()

                # Send OTP via email
                send_otp_via_email(email, otp)
                st.session_state['otp'] = otp
                st.session_state['email'] = email
            else:
                st.error('Invalid email or password.')

    # OTP verification
    if 'otp' in st.session_state:
        st.subheader('OTP Verification')
        entered_otp = st.text_input('Enter OTP')

        if st.button('Verify OTP'):
            if entered_otp == st.session_state['otp']:
                st.success('Login successful!')
                st.session_state.authenticated = True
                del st.session_state['otp']
                # st.experimental_rerun()
            else:
                st.error('Invalid OTP.')

# Main app logic
if st.session_state.authenticated:
    main_app()
else:
    login_page()
